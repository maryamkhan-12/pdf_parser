###################################################################################################################################
                                                    #Import Libraries
###################################################################################################################################
import io
import os
import re
import json
import shutil
import requests
from tqdm import tqdm
from PIL import Image
from pprint import pprint
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import List, Dict
from pydantic import BaseModel
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from langchain_groq import ChatGroq
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from langchain_core.prompts import PromptTemplate

###################################################################################################################################
                                                    #Create FastAPI
###################################################################################################################################


app = FastAPI()


###################################################################################################################################
# Define Models for Request Payloads
###################################################################################################################################

class ImagePromptRequest(BaseModel):
    blog_post_content: str
    previous_image_prompts: str

###################################################################################################################################
                                                    #Initilize LLM
###################################################################################################################################


def get_llm():
    """
    Returns the language model instance (LLM) using ChatGroq API.
    The LLM used is Llama 3.1 with a versatile 70 billion parameters model.

    Returns:
        llm (ChatGroq): An instance of the ChatGroq LLM.
    """
    llm = ChatGroq(
        model="llama-3.2-90b-text-preview",
        temperature=0.3,
        max_tokens=1024,
        api_key='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'
    )
    return llm

llm = get_llm()

###################################################################################################################################
                                                    #SEO
###################################################################################################################################

# Models for request payload
class blog_request(BaseModel):
    TypeOf : str
    target_audience: str
    tone: str
    point_of_view: str
    target_country: str
    keywords: List[str]

# Fetch Google results and format for SEO
def fetch_google_results(keywords: List[str], target_country: str) -> List[str]:
    username = 'madii_zMvf6'
    password = 'Momimaad_123'
    all_results_dict = {}

    # Fetch results for each keyword
    for keyword in keywords:
        payload = {
            'source': 'google_search',
            'query': keyword,
            'domain': 'com',
            'geo_location': target_country,
            'locale': 'en-us',
            'start_page': '1',
            'pages': '1',
            'context': [{'key': 'filter', 'value': 1}, {'key': 'results_language', 'value': 'en'}]
        }
        try:
            response = requests.post(
                'https://realtime.oxylabs.io/v1/queries',
                auth=(username, password),
                json=payload
            )
            response.raise_for_status()
            all_results_dict[keyword] = response.json()
        except requests.RequestException as e:
            raise HTTPException(status_code=500, detail=f"Error for '{keyword}': {str(e)}")

    # Format search results for use in title generation
    search_results = [
        f"Position {pos}: {result['Title']}"
        for keyword, results in all_results_dict.items()
        for pos, result in results.get('organic', {}).items()
    ]
    return search_results


###################################################################################################################################
                                                    # Generate Title for the BLOG Post
###################################################################################################################################


def generate_blog_title(keywords: List[str], search_results: List[str], category: str,blog_request: blog_request) -> str:

    """
    Generates an SEO-optimized title for a blog post using keywords and search results.

    Args:
        keywords (list): List of keywords provided by the user.
        all_results_dict (dict): Dictionary with search result data including position and titles.
        category (str): Category for the blog post (e.g., 'Parenting Stages', 'Self-Care for Moms').

    Returns:
        str: SEO-optimized blog title.
    """

    # Prompt template for generating a blog title
    prompt_template = """
    You are a skilled content creator with expertise in SEO and keyword optimization. Your task is to generate a compelling, SEO-optimized title for a blog post. You have to use following keywords in the title.

    Blog Post Category: {category}
    Keywords including position and title: {keywords}


    Instructions:
    - Write this Blog according {Type} type of Blog.
    - Use the provided keywords and category to create a unique, catchy, and SEO-friendly title.
    - The title should be concise, engaging, and reflect the topic in a way that attracts readers.
    - Aim for a title that is under 60 characters if possible, without sacrificing clarity.
    - Capture the essence of the topic, catering specifically to the interests of readers in the category of {category}.
    - Ensure the title has a positive, inviting tone appropriate for readers interested in {category}.
    - Answer should be a single title and do not include anything in the answer
    Based on the information provided, generate a title that fulfills these requirements.
    """

    # Format the prompt with the provided category, keywords, and search results
    prompt = prompt_template.format(
        Type=blog_request.TypeOf,
        category=category,
        keywords=", ".join(keywords),
        search_results="\n".join(search_results)
    )

    # Assuming llm is an instance of an LLM like GPT
    response = llm.invoke(prompt)
    return response.content

###################################################################################################################################
                                                    # Generate Sub Headings for the BLOG Post
###################################################################################################################################

def generate_blog_subheadings(title: str, search_results : list, seleted_catagory:str, blog_request: blog_request) -> List[str]:
    """
    Generates suggested subheadings for a blog post based on the provided details.

    Args:
        title (str): The main title of the blog post.
        seo_keywords (list): List of keywords to focus on in the content.
        target_audience (str): Description of the target audience.
        tone (str): The tone of the blog post (e.g., friendly, informative).
        point_of_view (str): The narrative perspective (e.g., first-person, third-person).
        target_country (str): The primary country where the audience is located.

    Returns:
        list: A list of suggested subheadings.
    """
    # Sample prompt to guide the subheadings
    prompt_template = """
    You are a content strategist tasked with creating SEO-friendly subheadings for a blog post.
    Generate clear, engaging, and informative subheadings to support the main title and resonate
    with the target audience.

    Title: {title}
    Seleted Catagory:{seleted_catagory}
    SEO Keywords: {search_results}
    Target Audience: {target_audience}
    Tone: {tone}
    Point of View: {point_of_view}
    Target Country: {target_country}

    Instructions:
    - Generate 1 subheadings that cover key aspects of the topic.
    - Ensure subheadings reflect the tone and point of view, and relate closely to the title and keywords.
    - Consider including helpful tips, important facts, and region-specific insights if relevant.
    - Keep subheadings concise and engaging.
    - Write only subheadings, do not include anything in the answer
    - Ensure the subheadings flows logically, with smooth transitions between subheadings.

    Based on these inputs, suggest subheadings for the blog post.
    """

    # Format the prompt with provided inputs
    prompt = prompt_template.format(
        title=title,
        seleted_catagory=seleted_catagory,
        search_results=", ".join(search_results),
        target_audience=blog_request.target_audience,
        tone=blog_request.tone,
        point_of_view=blog_request.point_of_view,
        target_country=blog_request.target_country
    
    )

    # Generate subheadings using the language model
    response = llm.invoke(prompt)
    suggested_subheadings = response.content.split("\n")

    # Return the list of suggested subheadings
    return [subheading.strip() for subheading in suggested_subheadings if subheading.strip()]

###################################################################################################################################
                                                    # Generate content for every sub headings
###################################################################################################################################

# Function to create content for a single subheading
def BlogPostPromptSingleSubheading(title: str, current_subheading: str, blog_request: blog_request, search_results: List[str], category: str, previous_content: str) -> str:
    # Template for generating content under each subheading
    prompt_template = """
    You are an advanced language model designed to help create professional blog posts
    specifically in the following category: {category}. Your task is to draft a well-structured, unique,
    factual, and engaging blog post section under each subheading that provides value to the readers.
    Use the provided keywords in each section to make it SEO-friendly. The section should be max 2-3 paragraphs.

    Blog Post Category: {category}
    Title: {title}
    Target Audience: {target_audience}
    Tone: {tone}
    Point of View: {point_of_view}
    Target Country: {target_country}

    Previous Content:
    {previous_content}

    Subheading: {current_subheading}
    Instructions:
    - Write a detailed, engaging section under this subheading that flows well with the previous content.
    - Use a tone that resonates with the target audience (e.g., reassuring, informative, or friendly).
    - Include real-time data, research findings, or expert opinions if relevant.
    - Ensure smooth transitions and consistency in content flow.
    - Incorporate practical tips, advice, or answers to potential questions related to this subheading.
    - Avoid irrelevant information; stay focused on the topic.
    - Write content that naturally leads into the next subheading.
    - Do not add Conclusions and FAQs in the blog
    - do not use ** with bullet points Use numbers for the bullet points only
    - do not use number with sub headinngs Use ** with sub headinngs
    - Only use ** for headings and sub headings only
    Based on this input, write content for the current subheading.
    """

    # Format the prompt with necessary parameters
    prompt = prompt_template.format(
        title=title,
        category=category,
        target_audience=blog_request.target_audience,
        tone=blog_request.tone,
        keywords=", ".join(search_results),
        point_of_view=blog_request.point_of_view,
        target_country=blog_request.target_country,
        previous_content=previous_content,
        current_subheading=current_subheading
    )

    # Generate content for the current subheading using the language model
    response = llm.invoke(prompt)
    content=response.content
    
    return content

###################################################################################################################################
                                                    # Format the document
###################################################################################################################################

def format_content(document, content: str):
    # Regular expression patterns to detect subheadings and bullet points
    subheading_pattern = r"\*\*(.*?)\*\*"
    bullet_point_pattern = r"^\s*•\s*\*\*\s*(.*?)\s*\*\*"

    # Split content by lines to process each line individually
    lines = content.split("\n")
    for line in lines:
        if re.match(subheading_pattern, line):
            # If line is a subheading, add it as a Heading level 2
            subheading_text = re.sub(r"\*\*", "", line).strip()
            document.add_heading(subheading_text, level=2)
        elif re.match(bullet_point_pattern, line):
            # If line is a bullet point, add it with bold text
            bullet_text = re.sub(r"^\s*•\s*\*\*", "", line).strip()
            p = document.add_paragraph(style='List Bullet')
            run = p.add_run(bullet_text)
            run.bold = True
        else:
            # Add other content as a paragraph with justified alignment
            p = document.add_paragraph(line.strip())
            p.alignment = 3  # Justify alignment

        # Set text color to black for consistency
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black


###################################################################################################################################
                                                    # Decide to generate image or not
###################################################################################################################################

def decide_to_generate_image(content: str, i:int) -> bool:
    max=3
    # Template for generating content under each subheading
    prompt_template = '''
    You are an advanced language model tasked with deciding if an image should be generated based on the provided blog post. Analyze the blog content and respond with "Yes" or "No" only. Generate an image only if the content is rich, high-quality, and would benefit from it. Generate a maximum of 3 images: if {i} > {max}, respond with "No".

    Blog post:
    {blog_post}

    Output:
    Yes or No
    '''

    prompt=prompt_template.format(blog_post=content,i=i,max=max)
    response = llm.invoke(prompt)

    should_generate_image = response.content
    return should_generate_image


###################################################################################################################################
                                                    # Generate image prompt
###################################################################################################################################

def generate_image_prompt(content: str, previous_prompts: str) -> str:
    prompt_template = """
    You are a creative assistant tasked with generating precise and captivating visual representation prompts
    for a blog post related to child-care topics.

    Blog Post Draft:
    {blog_post_content}

    Previous Image Prompts:
    {previous_image_prompts}

    Instructions:
    - Carefully analyze the blog post content to identify the most important and impactful sections, key points,
      or themes that would benefit from visual representation.
    - Focus on transforming these key points into well-structured and concise prompts for generating unique and
      accurate images that align with the blog content.
    - Make sure each prompt clearly conveys the visual idea, including essential details like style, setting,
      objects, or mood, while keeping it to the point.
    - Ensure prompts are diverse and cover various aspects of the blog to provide a comprehensive visual
      experience.
    - Avoid writing anything except the prompts themselves.
    - Limit the number of prompts to 1
    - Do not add prompts in which images involve text in it.
    - New prompt should be different from the previous prompt.
    - Write only prompts, do not include anything in the answer
    """

    # Format the prompt with the blog post draft
    prompt = prompt_template.format(blog_post_content=content,previous_image_prompts=previous_prompts)

    # Invoke the LLM to generate the image prompts
    response = llm.invoke(prompt)

    return response.content  # Extract prompts

###################################################################################################################################
                                                    # Generate image 
###################################################################################################################################

# Generate image using an external API
def generate_image(prompt: str):

    # Fetch the API token from environment variables
    API_TOKEN = os.getenv("HF_API_TOKEN")

    if not API_TOKEN:
        raise ValueError("API token is missing. Please set HF_API_TOKEN.")


    # Hugging Face API setup for image generation
    API_URL = "https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-dev"
    headers = {"Authorization": "Bearer {API_TOKEN}"}

    # Function to generate an image using Hugging Face's FLUX model
    payload = {"inputs": prompt}
    response = requests.post(API_URL, headers=headers, json=payload)

    if response.status_code == 200:
        image_bytes = response.content
        image = Image.open(io.BytesIO(image_bytes))
        return image
    else:
        print("Failed to generate image.")
        return None


###################################################################################################################################
                                                    # Selecte Catagory
###################################################################################################################################

def selected_category(category: dict, search_results: list) -> str:
    # Define the prompt template
    prompt_template = """
    Based on the given search results, select the most appropriate category for the blog post.
    
    Available Categories: {categories}
    
    Search Results: 
    {search_results}

    Carefully analyze the keywords and context in the search results to choose the best category. 
    Please respond only with the most relevant category name.
    """

    # Format the prompt with the category and search results
    prompt = prompt_template.format(categories=", ".join(category.keys()), search_results="\n".join(search_results))

    # Invoke the LLM to generate the selected category
    response = llm.invoke(prompt)

    return response.content.strip()  # Extract the selected category


###################################################################################################################################
                                                    # SEO Blog Post
###################################################################################################################################


# Fetch Google results for a specific site with dynamic keywords
def fetch_google_results_for_site(keywords: List[str]) -> List[Dict[str, int]]:
    username = 'madii_zMvf6'
    password = 'Momimaad_123'

    # Join keywords with '+' to form the query parameter
    query_string = "+".join(keywords)
    search_url = f"https://www.google.com/search?q=site:marcusmcdonnell.com+{query_string}"

    # Set up payload with direct URL search on Google
    payload = {
        'source': 'google',
        'url': search_url,
        'parse': True  # Enabling parsed response to get structured data
    }

    try:
        # Send initial POST request
        response = requests.post(
            'https://realtime.oxylabs.io/v1/queries',
            auth=(username, password),
            json=payload
        )
        response.raise_for_status()

        # Get the initial response
        full_response = response.json()

        # Navigate to 'results -> content -> results -> organic'
        filtered_results = []
        if full_response.get('results'):
            for result in full_response['results']:
                organic_results = result.get('content', {}).get('results', {}).get('organic', [])
                if isinstance(organic_results, list):
                    filtered_results.extend(
                        {"title": item.get("title"), "url": item.get("url"), "pos": item.get("pos")}
                        for item in organic_results
                        if "title" in item and "url" in item and "pos" in item
                    )
                else:
                    print("Expected 'organic' results to be a list but found something else.")
        else:
            print("No 'results' key found in the response.")

        return filtered_results

    except requests.RequestException as e:
        print(f"Error fetching results: {e}")
        return []

def generate_linkages(blog_post: str, search_results: list, keywords: List[str]) -> dict:
    """
    Generate external and internal linkages for a blog post using web search results.
    Uses the top 3 search results for external linkages and LLM to create the link section.
    
    Parameters:
    blog_post (str): The content of the blog post.
    search_results (list): The search results to analyze.

    Returns:
    dict: A dictionary with 'external_links' and 'internal_links' as keys.
    """

    Internal_search_results=fetch_google_results_for_site(keywords)

    # Define the prompt template
    prompt_template = """
    Based on the given blog post and search results, generate relevant external and internal links.

    Blog Post:
    {blog_post}

    Use the top 3 search results for external link suggestions, considering their relevance and quality.
    Also, suggest internal links that might help the reader based on the blog post's content.

    External Links:
    For each link, summarize the content and explain its relevance to the blog post. Limit to 3 high-quality links. If links are not avaliable then don't add them
    
    Internal Links:
    Suggest appropriate internal links based on potential topics within the blog post that readers may find useful. Limit to 3 high-quality links. If in Internal Links Results are not avaliable then don't add them in the internal links

    External Links Results:
    {search_results}

    Internal Links Results:
    {Internal_search_results}

    Output:
    External Links: A list of 3 high-quality links with brief explanations.
    Internal Links: A list of internal topics or pages with brief explanations.
    """

    # Format the prompt with the blog post and top search results
    prompt = prompt_template.format(blog_post=blog_post, search_results=search_results, Internal_search_results=Internal_search_results)

    # Invoke the LLM to generate the link suggestions
    response = llm.invoke(prompt)
    
    return response.content.strip()  # Adjust based on LLM output structure


###################################################################################################################################
                                                    # Generate Blog Post
###################################################################################################################################

# MAIN PIPELINE: Create blog with title, subheadings, content, and images
@app.post("/blogs/pipeline/", response_model=dict)
def create_blog_pipeline(blog_request: blog_request):

    # Define the file path for the generated blog post
    file_path = "Generated_Blog_Post.docx"

    # Check if the file exists, and if so, remove it
    if os.path.exists(file_path):
        os.remove(file_path)


    if not os.path.exists("pic"):
        os.makedirs("pic")

    category = {
    "Parenting Stages": [
        "Baby & Toddler Years",
        "Preschool & Early Childhood",
        "Big Kids (6–12 Years)",
        "Tweens & Teens",
        "Newborn Care",
        "Parenting After Baby #2 (or #3!)"
    ],
    "Everyday Life with Kids": [
        "Daily Routines & Organization",
        "Mealtime & Nutrition",
        "Playtime & Activities",
        "Sleep Schedules & Tips",
        "Family Time Ideas",
        "Special Occasions & Holidays"
    ],
    "Self-Care for Moms": [
        "Health & Wellness",
        "Mental Health & Stress Relief",
        "Beauty & Self-Care Tips",
        "Hobbies & “Me Time”",
        "Personal Growth & Goal Setting"
    ],
    "Parenting Tips & Tricks": [
        "Time-Saving Hacks",
        "Budgeting for Families",
        "Quick Cleaning Tips",
        "Home Organization with Kids",
        "School & Homework Help",
        "Tech Tools for Parenting"
    ],
    "Mom Life (Humor & Reality)": [
        "Honest Mom Moments",
        "Laughs & Parenting Memes",
        "Confessions & Fails",
        "Real Life, Real Moms",
        "Quotes & Relatable Stories"
    ],
    "Parenting Styles & Philosophies": [
        "Gentle Parenting & Positive Discipline",
        "Attachment Parenting",
        "Raising Independent Kids",
        "Balancing Structure & Freedom",
        "Parenting Trends & Research"
    ],
    "Relationships & Family Dynamics": [
        "Co-Parenting & Communication",
        "Sibling Relationships",
        "Family Bonding Activities",
        "Grandparents & Extended Family",
        "Blended Families & Step-Parenting"
    ],
    "Learning & Development": [
        "Early Childhood Education",
        "Fun Learning at Home",
        "Language & Social Skills",
        "Emotional Development",
        "Physical & Motor Skills"
    ],
    "Health & Wellness": [
        "Child Nutrition & Health",
        "Exercise & Play for Kids",
        "Sleep Health",
        "Pediatric Check-Ups",
        "Common Illnesses & Remedies",
        "Childproofing & Safety"
    ],
    "Mompreneurs & Working Moms": [
        "Balancing Work & Family",
        "Remote Work Tips",
        "Side Hustles & Passions",
        "Time Management for Busy Moms",
        "Self-Care for Working Moms"
    ],
    "School & Education": [
        "Preschool & Daycare Choices",
        "School Readiness & Transitions",
        "Homework & Study Skills",
        "Navigating School Friendships",
        "Involvement in School Activities"
    ],
    "Lifestyle & Home": [
        "Home Décor for Families",
        "Sustainable & Eco-Friendly Choices",
        "Family Finances & Budgeting",
        "Travel & Family Adventures",
        "Pets & Kids"
    ],
    "Parenting Challenges": [
        "Tantrums & Discipline",
        "Social Media & Screen Time",
        "Bullying & Peer Pressure",
        "Dealing with Picky Eaters",
        "Navigating Kids’ Fears"
    ],
    "Creative & Fun Ideas": [
        "DIY Projects for Kids",
        "Kid-Friendly Crafts",
        "Fun Recipes & Snacks",
        "Family Games & Activities",
        "Fun Celebrations & Birthdays"
    ],
    "Modern Parenting Topics": [
        "Raising Kids in a Digital World",
        "Cultural & Diversity Awareness",
        "Gender-Neutral Parenting",
        "Parenting and Social Media"
    ],
    "The Wild World of Parenting": [
        "Tiny Bosses: Life with Toddlers",
        "Kid Logic: Decoding the Mind of a Child",
        "Growing Up Fast: Navigating the Tween Years"
    ],
    "The Mom Zone": [
        "Surviving the Madness, One Coffee at a Time",
        "Keeping It Real: The Beautiful Mess of Mom Life",
        "Dear Diary: Honest Mom Confessions"
    ],
    "Mastering the Art of Family Chaos": [
        "Organized Chaos: Tips for Running a Household",
        "Barely Hanging On: Parenting Hacks for the Real World",
        "Kid-Proof Your Life (If That’s Even Possible)"
    ],
    "Mom Self-Care, Simplified": [
        "Time for You: Self-Care for Busy Moms",
        "Staying Sane (Mostly) with Self-Care on a Budget",
        "Love Yourself: From Self-Care to Self-Love"
    ],
    "Making Memories, Keeping Your Sanity": [
        "Everyday Magic: Fun Family Traditions",
        "Making the Ordinary Extraordinary",
        "The Cool Mom’s Guide to Family Fun"
    ],
    "Mom Hacks & Life-Saving Tricks": [
        "Shortcuts for Sanity: Clever Parenting Hacks",
        "The No-Fuss Guide to Getting Stuff Done",
        "Mom Brain Solutions: Little Tricks for Big Wins"
    ],
    "When Kids Are…Kids!": [
        "Real Talk: Tantrums, Tears & Tiny Attitudes",
        "Kid Quirks: The Weird, Wonderful World of Children",
        "Mini People, Mega Emotions"
    ],
    "Relationships and Realities": [
        "It Takes Two: Parenting Together (Even When You Don’t Agree)",
        "Keeping Love Alive Amid the Chaos",
        "Keeping the Family Peace, One Day at a Time"
    ],
    "The School Scene": [
        "Homework Without the Headache",
        "From Preschool to Preteen Drama: Surviving School Years",
        "Winning at School (Even If They Don’t Love It)"
    ],
    "Digital World for Digital Kids": [
        "Screen Time vs. Play Time: Finding the Balance",
        "Raising Tech-Savvy Kids in a Tech-Obsessed World",
        "Social Media & Selfies: Teaching Digital Smarts"
    ],
    "Raising the Next Generation": [
        "The Kindness Project: Raising Empathetic Kids",
        "How to Raise Future World-Changers",
        "The Power of Yes and No: Teaching Choices"
    ],
    "Healthier, Happier Families": [
        "Making Meals Easy & Fun (Yes, Really!)",
        "Health Hacks for Kids Who Hate Veggies",
        "Small Habits for Big Health Wins"
    ],
    "The Organized Chaos Hub": [
        "Declutter Like a Pro (Yes, Even with Kids)",
        "Home Hacks for the Ultimate Kid-Friendly Space",
        "Mastering the Family Schedule"
    ],
    "Funny Mom Survival Kit": [
        "Parenting Memes You’ll Feel in Your Soul",
        "Surviving Kids’ Parties with Style",
        "Confessions of a Bedtime Warrior"
    ],
    "Big Dreams & Little Goals": [
        "Goal-Getting for Moms Who Do It All",
        "Dare to Dream Big (Even If You’re Tired)",
        "Mom Goals: From ‘Just Survive’ to ‘Thrive’"
    ],
    "For the Love of Learning": [
        "Learning Through Play: Fun Ideas for Little Learners",
        "Home Learning Hacks for Smart Kids",
        "Raising Curious Kids: Sparking Little Imaginations"
    ],
    "Tales from the Trenches": [
        "Stories from the Wild World of Parenting",
        "Lessons Learned from the Chaos",
        "Hilarious Mom Stories You’ll Never Believe"
    ],
    "Adventures Big and Small": [
        "Tiny Adventures: Fun for Kids of All Ages",
        "Family Vacations & Kid-Friendly Getaways",
        "Staycations That Feel Like the Real Deal"
    ],
    "The Support Network": [
        "For the Love of Moms: Support & Community",
        "Village of Moms: Finding Your Support Circle",
        "Surviving & Thriving Together"
    ],
    "Creative Kids Zone": [
        "Arts & Crafts that Won’t Break the Bank",
        "Imagination Station: Encouraging Creative Play",
        "Rainy Day Fun: Indoor Ideas for Any Weather"
    ]
}
    print('SEO Searching')
    # Step 1: Fetch Google results
    search_results = fetch_google_results(blog_request.keywords, blog_request.target_country)
    
    print('Selecting Catagory for blog post')
    selected_cat=selected_category(category , search_results)

    print('Generating Title for blog post')
    previous_image_prompts = ''
    blog_content = ""
    document = Document()

    # Step 2: Generate Blog Title
    title = generate_blog_title(blog_request.keywords, search_results, selected_cat,blog_request)
    image_prompt = generate_image_prompt(title, previous_image_prompts)
    image = generate_image(image_prompt)
    image_path = f"pic/image.png"
    image.save(image_path)
    document.add_heading(title, 0)
    document.add_picture(image_path, width=Inches(6), height=Inches(6))


    print('Generating Subheadings for blog post')
    # Step 3: Generate Blog Subheadings
    subheadings = generate_blog_subheadings(title,search_results,selected_cat, blog_request)

    # Step 4: Loop through each subheading with progress tracking
    for i, subheading in enumerate(tqdm(subheadings, desc="Processing subheadings")):
        content = BlogPostPromptSingleSubheading(title, subheading, blog_request, search_results, selected_cat, blog_content)
        blog_content += f"\n\n{subheading}\n{content}"

        # Format and add content to the document
        format_content(document, content)

        # Step 5: Decide if an image should be generated
        if decide_to_generate_image(content,i):
            # Generate image prompt
            image_prompt = generate_image_prompt(content, previous_image_prompts)
            previous_image_prompts += image_prompt + " , "

            # Generate and save image
            image = generate_image(image_prompt)
            if image:
                image_path = f"pic/image_{i}.png"
                os.makedirs("pic", exist_ok=True)
                image.save(image_path)

                # Add the image to the document
                document.add_picture(image_path, width=Inches(6), height=Inches(6))

    
    blog_post=generate_linkages(blog_content,search_results,blog_request.keywords)

    # Format and add content to the document
    format_content(document, blog_post)

    # Step 6: Save the document
    document.save("Generated_Blog_Post.docx")

    # Step 7: Clean up images
    shutil.rmtree("pic", ignore_errors=True)
    return FileResponse(path="Generated_Blog_Post.docx", filename="Generated_Blog_Post.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
